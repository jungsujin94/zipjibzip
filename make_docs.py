"""
make_docs.py — 카드뉴스 자동화 시스템 설명 문서 + 콘텐츠 스크립트 생성
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ── 헬퍼 ──────────────────────────────────────────────────────────────────────

def add_heading(doc, text, level=1, color=None):
    p = doc.add_heading(text, level=level)
    if color:
        for run in p.runs:
            run.font.color.rgb = RGBColor(*color)
    return p

def add_body(doc, text, bold=False, italic=False, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    return p

def add_bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.font.bold = True
        run.font.size = Pt(11)
        p.add_run(text).font.size = Pt(11)
    else:
        run = p.add_run(text)
        run.font.size = Pt(11)
    return p

def add_numbered(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style='List Number')
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.font.bold = True
        run.font.size = Pt(11)
        p.add_run(text).font.size = Pt(11)
    else:
        p.add_run(text).font.size = Pt(11)
    return p

def shade_paragraph(p, hex_color="E8F0FE"):
    """단락 배경 음영 추가"""
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    pPr.append(shd)

# ── 문서 1: 시스템 설명서 ────────────────────────────────────────────────────

def make_explainer():
    doc = Document()

    # 기본 폰트
    style = doc.styles['Normal']
    style.font.name = 'Malgun Gothic'
    style.font.size = Pt(11)

    # 제목
    title = doc.add_heading('카드뉴스 자동화 시스템', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)

    sub = doc.add_paragraph('Claude AI + Python으로 만든 인스타그램 카드뉴스 자동 생성 파이프라인')
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in sub.runs:
        run.font.italic = True
        run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
        run.font.size = Pt(12)

    doc.add_paragraph()

    # ── 개요 ──
    add_heading(doc, '개요', level=1, color=(0x1A, 0x1A, 0x1A))
    add_body(doc,
        '오늘의집 또는 쿠팡 제품 URL 하나를 입력하면, AI가 자동으로 제품을 분석하고 '
        '인스타그램용 마케팅 카드뉴스 6장을 PNG로 생성합니다. '
        '업로드 승인 시 웹 카탈로그(index.html)까지 자동으로 업데이트됩니다.'
    )
    doc.add_paragraph()

    # ── 전체 흐름 ──
    add_heading(doc, '전체 파이프라인 흐름', level=1, color=(0x1A, 0x1A, 0x1A))

    steps = [
        ('STEP 1 — URL 입력',
         '사용자가 오늘의집 또는 쿠팡 제품 페이지 URL을 Claude Code 채팅에 붙여넣습니다. '
         '오zip.me 같은 단축 URL도 자동으로 원본 URL로 해석합니다.',
         [
             '지원 사이트: ohou.se (오늘의집), coupang.com (쿠팡)',
             '단축 URL (ozip.me, bit.ly 등) 자동 해석 및 추적 파라미터 제거',
         ]),
        ('STEP 2 — 제품 스크래핑 (scraper.js)',
         'Node.js + Playwright를 이용해 제품 페이지에 실제 브라우저로 접속하여 데이터를 수집합니다.',
         [
             '수집 항목: 제품명, 가격, 할인율, 브랜드, 제품 이미지 (최대 8장), 실제 구매자 리뷰, 평점',
             '봇 차단 우회를 위해 실제 브라우저 렌더링 방식 사용',
             '배송/설치 안내 이미지는 자동 필터링',
         ]),
        ('STEP 3 — AI 콘텐츠 분석 (extractor.py)',
         'Claude Sonnet API를 호출해 수집된 제품 데이터를 마케팅 카드뉴스 6장 구조로 분석합니다.',
         [
             'HOOK 카드: 스크롤을 멈추게 하는 공감 문구 (페인포인트 직격 / 상황 묘사 / 질문형 / 수치 팩트)',
             'COMBO 카드: 소비자 고민 3가지 + 구매 시 체크포인트 3가지',
             'REVIEW 카드: 실제 스크래핑된 리뷰를 40자 이내로 자연스럽게 요약',
             'STAT 카드: 제품 카테고리와 연관된 놀라운 통계/수치 (공유 유도)',
             'SOLUTION 카드: 제품 풀네임, 브랜드, 실제 가격, 핵심 USP',
             'CTA 카드: 팔로우 및 구매 유도',
             '인스타그램 캡션 + 해시태그 5개 자동 생성',
         ]),
        ('STEP 4 — 카드뉴스 렌더링 (card_renderer.py)',
         'Python Pillow 라이브러리로 1080×1080px PNG 이미지 6장을 자동 생성합니다.',
         [
             '각 카드마다 최적 레이아웃으로 텍스트·이미지 자동 배치',
             '제품 이미지 배경 자동 제거 (birefnet AI 모델)',
             '여러 제품 이미지 중 가장 선명한 단독 컷 자동 선별',
             '출력 위치: output/ 폴더',
         ]),
        ('STEP 5 — 업로드 승인 (yes / no / hold)',
         '생성된 6장 카드를 확인한 후 세 가지 명령으로 처리합니다.',
         [
             'yes → 카드 1번 이미지를 products/ 폴더에 저장 + products.json에 제품 정보 추가/업데이트',
             'no  → output/ 폴더의 이번 생성 파일 전체 삭제',
             'hold → output/ 폴더에 그대로 유지 (나중에 재검토)',
         ]),
        ('STEP 6 — 카탈로그 자동 업데이트 (page.py)',
         '"yes" 응답 시 index.html 웹 카탈로그에 새 제품이 자동으로 추가됩니다.',
         [
             '카테고리 탭 자동 생성 (의자, 테이블, 조명 등)',
             '가격 낮은순 / 높은순 / 최신순 정렬 기능',
             '당일·전일 업로드 제품에 NEW 스티커 자동 표시',
             '오늘의집 · 쿠팡 구매 링크 버튼 자동 삽입',
         ]),
    ]

    for i, (title_text, desc, bullets) in enumerate(steps):
        # 스텝 제목
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(8)
        run = p.add_run(f'  {title_text}  ')
        run.font.bold = True
        run.font.size = Pt(13)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        colors = [
            (0x2B, 0x2D, 0x42),
            (0x3A, 0x86, 0xFF),
            (0x8B, 0x5C, 0xF6),
            (0x06, 0xB6, 0xD4),
            (0x10, 0xB9, 0x81),
            (0xF5, 0x9E, 0x0B),
        ]
        shade_paragraph(p, '{:02X}{:02X}{:02X}'.format(*colors[i]))

        add_body(doc, desc)
        for b in bullets:
            add_bullet(doc, b)
        doc.add_paragraph()

    # ── 기술 스택 ──
    add_heading(doc, '기술 스택', level=1, color=(0x1A, 0x1A, 0x1A))
    tech = [
        ('언어', 'Python 3, Node.js (JavaScript)'),
        ('AI 모델', 'Claude Sonnet (Anthropic API) — 콘텐츠 분석 및 생성'),
        ('스크래핑', 'Playwright (headless Chromium 브라우저 자동화)'),
        ('이미지 처리', 'Pillow (PNG 렌더링), rembg/birefnet (배경 제거)'),
        ('웹 카탈로그', 'Pure HTML/CSS/JS (프레임워크 없음, GitHub Pages 배포 가능)'),
        ('실행 방법', 'Claude Code CLI — /cardnews [URL] 명령어'),
    ]
    for label, val in tech:
        add_bullet(doc, f'{val}', bold_prefix=f'{label}: ')

    doc.add_paragraph()

    # ── 파일 구조 ──
    add_heading(doc, '주요 파일 구조', level=1, color=(0x1A, 0x1A, 0x1A))
    files = [
        ('pipeline.py', '전체 파이프라인 오케스트레이터 (메인 진입점)'),
        ('scraper.js', 'Playwright 기반 제품 페이지 스크래퍼'),
        ('extractor.py', 'Claude API 호출 및 6장 콘텐츠 JSON 생성'),
        ('card_renderer.py', 'Pillow 1080×1080 카드 PNG 렌더러'),
        ('page.py', 'products.json → index.html 카탈로그 생성'),
        ('products.json', '업로드 승인된 제품 목록 데이터베이스'),
        ('products/', '카드 1번 이미지 저장 폴더 (카탈로그 썸네일)'),
        ('output/', '생성된 카드뉴스 PNG 6장 + 캡션 txt 임시 저장'),
        ('index.html', '웹 제품 카탈로그 (자동 생성)'),
    ]
    for fname, desc in files:
        add_bullet(doc, f'  — {desc}', bold_prefix=fname)

    doc.save('C:/cardnews/카드뉴스_자동화_시스템_설명서.docx')
    print('✅ 설명서 저장: 카드뉴스_자동화_시스템_설명서.docx')


# ── 문서 2: 콘텐츠 스크립트 ─────────────────────────────────────────────────

def make_script():
    doc = Document()

    style = doc.styles['Normal']
    style.font.name = 'Malgun Gothic'
    style.font.size = Pt(11)

    title = doc.add_heading('콘텐츠 스크립트', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub = doc.add_paragraph('카드뉴스 자동화 시스템 소개 영상 / 릴스 낭독용')
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in sub.runs:
        run.font.italic = True
        run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    doc.add_paragraph()

    # 안내
    note = doc.add_paragraph()
    shade_paragraph(note, 'FFF3CD')
    note.add_run('📌 사용 안내').font.bold = True
    note.add_run(
        '   [ ] 안의 내용은 상황에 맞게 조절하세요. '
        '(예시) 표시는 실제 예시이며, 본인 제품으로 대체 가능합니다.'
    )

    doc.add_paragraph()

    sections = [
        ('🎬 오프닝 (0~15초)', '후킹 — 시청자 눈길 잡기', [
            ('',
             '제품 하나 고르는 데 몇 시간 걸리세요?\n'
             '저는 링크 하나 붙여넣는 데 30초면 됩니다.',
             '(짧게 끊고, 잠깐 멈춤)'),
            ('',
             '오늘은 제가 Claude AI랑 직접 만든\n'
             '카드뉴스 자동화 시스템을 소개해 드릴게요.',
             None),
        ]),
        ('📌 시스템 소개 (15~45초)', '무엇인지 한 문장으로', [
            ('',
             '오늘의집이나 쿠팡 제품 링크를 하나 주면,\n'
             'AI가 제품을 분석하고\n'
             '인스타그램용 카드뉴스 6장을 자동으로 만들어 줍니다.',
             None),
            ('',
             '직접 타이핑할 필요 없고,\n'
             '이미지 편집 툴 켤 필요도 없어요.\n'
             '링크 → 카드뉴스, 끝입니다.',
             '(손가락 스냅 제스처)'),
        ]),
        ('⚙️ 작동 방식 (45초~2분)', '파이프라인 단계별 설명', [
            ('STEP 1 — 링크 입력',
             '제가 Claude Code라는 AI 코딩 도구에\n'
             '제품 링크를 그냥 붙여넣어요.\n'
             '오늘의집, 쿠팡, 단축 URL 모두 됩니다.',
             None),
            ('STEP 2 — 자동 스크래핑',
             '그러면 시스템이 직접 그 제품 페이지에 접속해서\n'
             '제품명, 가격, 사진, 실제 구매 후기까지\n'
             '전부 긁어와요. 저는 아무것도 안 해도 됩니다.',
             None),
            ('STEP 3 — AI 분석',
             '그다음 Claude AI가 그 데이터를 분석해서\n'
             '카드뉴스 6장 내용을 짜줘요.\n\n'
             '첫 번째 카드는 "스크롤 멈추는 후킹 문구",\n'
             '두 번째는 공감 포인트랑 구매 체크리스트,\n'
             '세 번째는 실제 후기 요약,\n'
             '네 번째는 공유하고 싶은 통계나 수치,\n'
             '다섯 번째는 제품 소개 + 가격,\n'
             '여섯 번째는 팔로우 유도 CTA.\n\n'
             '이게 전부 자동으로 만들어집니다.',
             '(카드 6장 화면 보여주기)'),
            ('STEP 4 — 이미지 렌더링',
             'AI가 짜준 내용을 바탕으로\n'
             '1080×1080 사이즈 PNG 이미지로 뽑아줘요.\n'
             '배경 제거도 자동이고,\n'
             '여러 제품 사진 중에 가장 깔끔한 컷을 골라서 넣어줍니다.',
             None),
            ('STEP 5 — 제 선택: yes / no / hold',
             '카드 6장이 완성되면 저한테 보여주고 물어봐요.\n'
             '"업로드할까요?" 라고요.\n\n'
             '"yes" 하면 — 제품 카탈로그에 자동으로 추가되고,\n'
             '"no" 하면 — 파일이 다 지워지고,\n'
             '"hold" 하면 — 일단 저장해두고 나중에 결정해요.\n\n'
             '이 한 단어로 전부 처리됩니다.',
             '(키보드 타이핑 화면)'),
        ]),
        ('🌐 카탈로그 업데이트 (2분~2분 30초)', '웹사이트 자동 반영', [
            ('',
             '"yes"를 치는 순간\n'
             '제 zipjibzip 제품 페이지에 자동으로 올라가요.\n\n'
             '오늘의집 구매 링크, 쿠팡 링크,\n'
             '카테고리 분류, 가격 정보까지 다요.\n\n'
             '오늘 올린 제품은 NEW 스티커도 붙고요.',
             '(index.html 화면 보여주기)'),
        ]),
        ('💬 마무리 + CTA (2분 30초~3분)', '마무리 및 시청자 참여 유도', [
            ('',
             '저는 Claude AI를 단순히 질문 답변 도구로 쓰는 게 아니라\n'
             '"같이 일하는 도구"로 쓰고 있어요.\n\n'
             '코딩을 잘 몰라도 괜찮아요.\n'
             '제가 원하는 게 뭔지 말할 수 있으면 충분합니다.',
             None),
            ('',
             '이런 자동화 시스템 더 궁금하신 분들은\n'
             '팔로우 해두시면 계속 보여드릴게요.\n\n'
             '어떤 걸 자동화하면 좋을 것 같으세요?\n'
             '댓글로 알려주시면 다음 콘텐츠 소재로 써볼게요. 🖤',
             '(카메라 보며 마무리)'),
        ]),
    ]

    for section_title, section_sub, lines in sections:
        # 섹션 헤더
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        r = p.add_run(f' {section_title} ')
        r.font.bold = True
        r.font.size = Pt(13)
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        shade_paragraph(p, '2B2D42')

        p2 = doc.add_paragraph()
        r2 = p2.add_run(section_sub)
        r2.font.italic = True
        r2.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
        r2.font.size = Pt(10)

        for label, script_text, direction in lines:
            if label:
                lp = doc.add_paragraph()
                lp.paragraph_format.space_before = Pt(6)
                lr = lp.add_run(f'[ {label} ]')
                lr.font.bold = True
                lr.font.size = Pt(11)
                lr.font.color.rgb = RGBColor(0x3A, 0x86, 0xFF)

            # 스크립트 본문
            sp = doc.add_paragraph()
            sp.paragraph_format.left_indent = Inches(0.3)
            sr = sp.add_run(script_text)
            sr.font.size = Pt(12)
            sr.font.name = 'Malgun Gothic'

            # 연출 지시
            if direction:
                dp = doc.add_paragraph()
                dp.paragraph_format.left_indent = Inches(0.3)
                dr = dp.add_run(f'🎥 {direction}')
                dr.font.italic = True
                dr.font.size = Pt(10)
                dr.font.color.rgb = RGBColor(0xAA, 0x66, 0x00)

            doc.add_paragraph()

    doc.save('C:/cardnews/카드뉴스_콘텐츠_스크립트.docx')
    print('✅ 스크립트 저장: 카드뉴스_콘텐츠_스크립트.docx')


if __name__ == '__main__':
    make_explainer()
    make_script()
    print('\n두 파일 모두 C:/cardnews/ 에 저장되었습니다.')
